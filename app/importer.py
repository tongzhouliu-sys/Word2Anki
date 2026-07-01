import re
import docx
from pathlib import Path

def clean_line(line: str) -> str | None:
    """
    Cleans a line of text: removes list numbering/bullets, strips whitespace,
    and validates that the line contains an English word or phrase.
    """
    line = line.strip()
    if not line:
        return None

    # Remove list prefixes like "1. ", "2) ", "- ", "* ", "• "
    line = re.sub(r'^\s*(\d+[\.\)]|[\-\*•])\s*', '', line).strip()
    
    # Must contain at least one English letter and only consist of valid English word/phrase characters
    if re.search(r'[a-zA-Z]', line) and re.match(r'^[a-zA-Z0-9\s\-\'\,\.\(\)\?\/’‘…\:\;\!\"“”]+$', line):
        return line
    return None

def extract_number_prefix(line: str) -> tuple[int | None, str]:
    """
    Extracts leading integer number from the line if it is a list prefix.
    Returns (parsed_int, cleaned_rest).
    """
    line = line.strip()
    # Match patterns like: "1. apple", "1) apple", "[1] apple", "1 - apple"
    # Ensure it's a prefix by checking for separating characters (dot, parenthesis, hyphen, brackets)
    match = re.match(r'^\s*(?:\[?(\d+)\]?[\.\)\-\s]+|\s*(\d+)[\.\)]\s*)(.*)$', line)
    if match:
        num_str = match.group(1) or match.group(2)
        rest = match.group(3).strip()
        # Ensure rest contains letters, otherwise it might just be a random number line
        if num_str and re.search(r'[a-zA-Z]', rest):
            return int(num_str), rest
            
    # Also handle simple "number space word" if it's clearly a list:
    # E.g. "1 apple" where there's a space after digits.
    match_space = re.match(r'^\s*(\d+)\s+([a-zA-Z].*)$', line)
    if match_space:
        num_str = match_space.group(1)
        rest = match_space.group(2).strip()
        return int(num_str), rest
        
    return None, line

def normalize_term(line: str) -> list[str]:
    """
    Returns the line as is. Slash expansion is disabled to respect the
    original 'one line = one card' format of the document.
    """
    return [line]

def extract_words_from_docx(file_path: str) -> tuple[list[str], int]:
    """
    Reads a .docx file line-by-line (from paragraphs and tables), cleans each line,
    normalizes slash-expressions, and extracts English words/phrases.
    
    If line number prefixes exist (e.g. "1. apple", "2. banana"), the list is sorted
    primarily by numbering order (highest priority), and secondarily by original
    document appearance order.
    
    Returns a tuple containing:
      - unique_terms: The list of sorted unique lowercased words/phrases.
      - raw_lines_count: The count of valid lines detected before expansion.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Word document not found at: {file_path}")

    doc = docx.Document(str(path))
    raw_lines = []
    
    # Extract from paragraphs
    for para in doc.paragraphs:
        if para.text:
            raw_lines.append(para.text)
            
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    raw_lines.append(cell.text)
                    
    raw_lines_count = 0
    collected_terms = []
    doc_index = 0
    for line in raw_lines:
        # Split by newline in case a cell or paragraph has multiple lines
        for sub_line in line.split('\n'):
            # Parse number prefix if present
            num, rest = extract_number_prefix(sub_line)
            cleaned = clean_line(rest)
            if cleaned:
                raw_lines_count += 1
                # Normalize and expand any slash terms
                normalized_terms = normalize_term(cleaned)
                for term in normalized_terms:
                    cleaned_term = clean_line(term)
                    if cleaned_term:
                        collected_terms.append({
                            "word": cleaned_term.lower(),
                            "num": num,
                            "index": doc_index
                        })
                        doc_index += 1
                
    # Sort primarily by numbering (num_val), and secondarily by original index
    def sort_key(item):
        num_val = item["num"] if item["num"] is not None else float('inf')
        return (num_val, item["index"])

    # Compute unique numbering count (number of distinct serial numbers detected)
    unique_numbers = {item["num"] for item in collected_terms if item["num"] is not None}
    numbered_count = len(unique_numbers)

    sorted_items = sorted(collected_terms, key=sort_key)
    all_terms = [item["word"] for item in sorted_items]
    return all_terms, raw_lines_count, numbered_count
